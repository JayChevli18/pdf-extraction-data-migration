"""Extract plain text from PDF and DOCX (local files)."""

from __future__ import annotations

import logging
from io import BytesIO
from pathlib import Path

import pdfplumber
from docx import Document

logger = logging.getLogger("profile_backend.text_extract")


def extract_text(path: Path) -> str:
    """Return UTF-8 text from a .pdf or .docx file."""
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _pdf_text(path)
    if suffix in (".docx",):
        return _docx_text(path)
    raise ValueError(f"Unsupported format: {suffix}")


def extract_text_bytes(suffix: str, data: bytes) -> str:
    """Return UTF-8 text from in-memory PDF/DOCX bytes."""
    s = suffix.lower()
    if s == ".pdf":
        return _pdf_text_filelike(BytesIO(data))
    if s == ".docx":
        return _docx_text_filelike(BytesIO(data))
    raise ValueError(f"Unsupported format: {suffix}")


def _pdf_text(path: Path) -> str:
    parts: list[str] = []
    try:
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                t = page.extract_text() or ""
                parts.append(t)
    except Exception as e:
        logger.exception("pdfplumber failed for %s: %s", path, e)
        raise
    return "\n".join(parts).strip()


def _pdf_text_filelike(f) -> str:
    parts: list[str] = []
    try:
        with pdfplumber.open(f) as pdf:
            for page in pdf.pages:
                t = page.extract_text() or ""
                parts.append(t)
    except Exception as e:
        logger.exception("pdfplumber failed for bytes: %s", e)
        raise
    return "\n".join(parts).strip()


def _docx_text(path: Path) -> str:
    try:
        doc = Document(str(path))
        return "\n".join(p.text for p in doc.paragraphs if p.text).strip()
    except Exception as e:
        logger.exception("python-docx failed for %s: %s", path, e)
        raise


def _docx_text_filelike(f) -> str:
    try:
        doc = Document(f)
        return "\n".join(p.text for p in doc.paragraphs if p.text).strip()
    except Exception as e:
        logger.exception("python-docx failed for bytes: %s", e)
        raise
