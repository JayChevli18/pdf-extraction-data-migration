"""Text extraction adapters for PDF/DOCX content."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

import pdfplumber
from docx import Document


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _pdf_text(path)
    if suffix == ".docx":
        return _docx_text(path)
    raise ValueError(f"Unsupported format: {suffix}")


def extract_text_bytes(suffix: str, data: bytes) -> str:
    suffix = suffix.lower()
    if suffix == ".pdf":
        return _pdf_text_filelike(BytesIO(data))
    if suffix == ".docx":
        return _docx_text_filelike(BytesIO(data))
    raise ValueError(f"Unsupported format: {suffix}")


def _pdf_text(path: Path) -> str:
    parts: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            parts.append(page.extract_text() or "")
    return "\n".join(parts).strip()


def _pdf_text_filelike(file_like) -> str:
    parts: list[str] = []
    with pdfplumber.open(file_like) as pdf:
        for page in pdf.pages:
            parts.append(page.extract_text() or "")
    return "\n".join(parts).strip()


def _docx_text(path: Path) -> str:
    doc = Document(str(path))
    return "\n".join(p.text for p in doc.paragraphs if p.text).strip()


def _docx_text_filelike(file_like) -> str:
    doc = Document(file_like)
    return "\n".join(p.text for p in doc.paragraphs if p.text).strip()
